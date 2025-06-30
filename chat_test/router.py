from fastapi import APIRouter, HTTPException, status, UploadFile,File,Form
from pydantic import BaseModel
#from views import get_answer,rag,comparar_documentos
#from PyPDF2 import PdfReader
#from langchain.prompts import PromptTemplate
from openai import AzureOpenAI


router_ = APIRouter()

class Questions(BaseModel):
    ask: str


import os
import re
current_dir = os.path.dirname(os.path.abspath(__file__))


client = AzureOpenAI(
    api_version=os.getenv("API_VERSION_4_1"),
    azure_endpoint=os.getenv("ENDPOINT_DANI"),
    api_key=os.getenv("API_KEY_DANI"),
)


async def read_md(file_path: str):
    doc_dir = os.path.normpath(os.path.join(current_dir,file_path))
    with open(file=doc_dir, mode="r", encoding="utf-8") as file:
        prompt = file.read()
        return prompt
    
async def procesar_prospecto(md_text):

    processed_text = re.sub(r'~~(.*?)~~', r'<span style="color: red;">~~\1~~</span>', md_text)

    processed_text = re.sub(r'\*\*(.*?)\*\*', r'<span style="color: green;">**\1**</span>', processed_text)
    return processed_text
    


async def chat_completation_doc(md_text: str):

    system_prompt = """"
                        Recibiras un texto en formato de Markdown.
                        Quitar del texto MD los elementos "**".
                        Quitar del texto MD los elementos tachados con "~~".
                        Quitar encabezados desde la segunda página en adelante.
                        Quitar los números de página.
                        La salida debe ser texto Markdown unicamente, sin introducción ni conclusión."""
    

    
    response = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": f"{system_prompt}",
            },
            {
                "role": "user",
                "content": f"Texto MD:{md_text}",
            }
        ],
        temperature=0.5,
        tools=[],
        top_p=1.0,
        model="gpt-4.1"
    )
    respuesta = response.choices[0].message.content
    return respuesta


async def chat_completation_doc(md_text: str):

    system_prompt = """"
                        Recibiras un texto en formato de Markdown.
                        Quitar del texto MD los elementos "**".
                        Quitar del texto MD los elementos tachados con "~~".
                        Quitar encabezados desde la segunda página en adelante.
                        Quitar los números de página.
                        La salida debe ser texto Markdown unicamente, sin introducción ni conclusión."""
    

    
    response = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": f"{system_prompt}",
            },
            {
                "role": "user",
                "content": f"Texto MD:{md_text}",
            }
        ],
        temperature=0.5,
        tools=[],
        top_p=1.0,
        model="gpt-4.1"
    )
    respuesta = response.choices[0].message.content
    return respuesta



async def chat_comparation(prompt: str):
    response = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": f"{prompt}",
            },
            {
                "role": "user",
                "content": "",
            }
        ],
        #max_tokens=4096,
        temperature=0.5,
        tools=[],
        top_p=1.0,
        model="gpt-4.1"
    )
    respuesta = response.choices[0].message.content
    return respuesta

from docx import Document 

async def create_docx(parragragh: str):
    doc = Document()  

    doc.add_heading('Reporte Limpio', level=1)  

    doc.add_paragraph(f'{parragragh}')

    doc.save('ejemplo_reporte_limpio.docx')

    print("Listo creado el documento docx")


 





    
    
@router_.get("/")
async def root():
    return {"message":"Bienvenido al Backend"}
    


from fastapi import UploadFile, File, APIRouter
from langchain.prompts import PromptTemplate
from views import TempleFileHannder
import asyncio  

@router_.post("/mini_rag")
async def comparar_docus(file_1: UploadFile = File(...), file_2: UploadFile = File):
    files_list = [file_1,file_2]

    temp_doc_pdf = None  
    temp_pdf = None  

    archivos_temp = []

    for file in files_list:
        if file.filename.lower().endswith(".docx"):
            temp_doc_pdf = await TempleFileHannder.create_temple_doc_to_pdf(file=file.file)
            archivos_temp.append(temp_doc_pdf)
        elif file.filename.lower().endswith(".pdf"):
            temp_pdf = await TempleFileHannder.create_templife_pdf(file=file.file)
            archivos_temp.append(temp_pdf)

    temp_md_1 = await TempleFileHannder.make_md_file(temp_path=temp_pdf)

    archivos_temp.append(temp_md_1)

    temp_md_2 = await TempleFileHannder.make_md_file(temp_path=temp_doc_pdf)

    archivos_temp.append(temp_md_2)

    prompt_template = PromptTemplate(
        input_variables=["md_isp","md_adium"],
        template= await read_md("promptSystem_5.md")
    )

    with open(file=temp_md_1, mode="r", encoding="utf-8") as file:
        md_data_1 = file.read()

    llm_clean_md = await chat_completation_doc(md_text=md_data_1)
    await create_docx(parragragh=llm_clean_md)

    md_colors = await procesar_prospecto(md_text=md_data_1)

    file_new = current_dir + "/archivo_prueba_colores.md"
    
    with open(file_new, mode="w", encoding="utf-8") as f:
            f.write(md_colors)

    with open(file=temp_md_2, mode="r", encoding="utf-8") as file:
        md_data_2 = file.read()


    final_prompt = prompt_template.format(
        md_isp=md_colors,
        md_adium=md_data_2,
    )

    await asyncio.sleep(0.2)  
    # Elimina todos los archivos temporales creados, si existen  
    for temp_path in archivos_temp:  
        if temp_path and os.path.exists(temp_path):  
            try:  
                os.remove(temp_path)  
            except Exception as e:  
                print(f"Error al borrar {temp_path}: {str(e)}")  

    
    response = await chat_comparation(prompt=final_prompt)

    file_new = current_dir + "/archivo_final.md"
    try:
        with open(file_new, mode="w", encoding="utf-8") as f:
            f.write(response)
            return {"response_comparation":"ARCHIVO MARKDOWN LISTO!"}
    except Exception as e:
        raise Exception(e)


    

    

# @router_.post("/document_int")
# async def comparar_docus():
#     # if not file_1:
#     #     raise HTTPException(status_code=403, detail="Los campos no pueden estar vacíos")

#     answer_chat = await doc_intelligence()
#     return answer_chat


# @router_.post("/doc_int")
# async def doc_int_route(file_1: UploadFile = File(...)):
#     lines = await analyze_document_with_azure(file_1)
#     return {"lines": lines}



# @router_.post("/chat_llm",status_code=status.HTTP_201_CREATED)
# async def conversacion_llm(question: Questions):
#     ask = question.ask
#     answer_chat = await get_answer(ask)
#     return{"respuesta":answer_chat}



# @router_.post("/chat_agentes",status_code=status.HTTP_201_CREATED)
# async def conversacion_llm(question: Questions):
#     ask = question.ask
#     answer_chat = await get_message(ask)
#     return{"respuesta":answer_chat}




# @router_.post("/upload_file")
# async def add_document(title_1: str = Form(...), file_1: UploadFile = File(...),title_2: str = Form(...), file_2: UploadFile = File(...)):
    
#     if not title_1 or not file_1:
#         raise HTTPException(status_code=403, detail="Los campos no pueden estar vacíos")
#     elif not title_2 or not file_2:
#         raise HTTPException(status_code=403, detail="Los campos no pueden estar vacíos")
    

#     reader = PdfReader(file_1.file)

#     text_1 = ""
#     for page in reader.pages:
#         text_1 += page.extract_text()


#     reader_2 = PdfReader(file_2.file)

#     text_2 = ""
#     for page_2 in reader_2.pages:
#         text_2 += page_2.extract_text()

#     prompt_summary = open_prompt("promptSystem.md")

#     prompt = PromptTemplate(
#         input_variables=["document_a", "document_b"],
#         template=prompt_summary,
#     )


#     final_prompt = prompt.format(document_a=text_1,document_b=text_2)

#     print(final_prompt)

#     answer_chat = await get_answer(system=final_prompt)
#     return answer_chat


